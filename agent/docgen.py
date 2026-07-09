"""
Renders the final ExecutionPlan (with drafted section content) into a
polished .docx using python-docx: title page, styled headings,
tables for structured/mock data, page numbers in the footer, and a
consistent typography theme.
"""

from __future__ import annotations
import os
import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .schemas import ExecutionPlan

ACCENT_COLOR = RGBColor(0x1F, 0x4E, 0x79)
OUTPUT_DIR = os.getenv("OUTPUT_DIR", "generated_docs")


def _set_base_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    heading1 = doc.styles["Heading 1"]
    heading1.font.color.rgb = ACCENT_COLOR
    heading1.font.size = Pt(18)


def _add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def _add_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Page ")
    _add_page_number_field(p)
    p.add_run(" | Generated autonomously by DraftPilot")


def _add_title_page(doc: Document, plan: ExecutionPlan, user_request: str):
    doc.add_paragraph().paragraph_format.space_before = Pt(60)
    title = doc.add_heading(plan.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = ACCENT_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(plan.document_type.replace("_", " ").title())
    r.italic = True
    r.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Prepared for: {plan.audience}\n").font.size = Pt(11)
    meta.add_run(f"Date: {datetime.date.today().strftime('%B %d, %Y')}").font.size = Pt(11)

    doc.add_page_break()


def _add_assumptions(doc: Document, plan: ExecutionPlan):
    doc.add_heading("Assumptions Made by the Agent", level=1)
    doc.add_paragraph(
        "Because the original request left some details unspecified, the agent "
        "made the following explicit assumptions in order to proceed autonomously:"
    )
    for a in plan.assumptions:
        doc.add_paragraph(a, style="List Bullet")


def _add_table(doc: Document, table_data: dict):
    headers = table_data["headers"]
    rows = table_data["rows"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    doc.add_paragraph()  # spacing after table


def _add_sections(doc: Document, plan: ExecutionPlan):
    for section in plan.sections:
        doc.add_heading(section.heading, level=1)
        if section.content:
            for para in section.content.split("\n"):
                para = para.strip()
                if not para:
                    continue
                # Defensive filter: if the model ignored the "no markdown tables"
                # instruction, drop pipe-delimited table rows/separators rather
                # than dumping raw "| a | b |" text into the Word document.
                if para.count("|") >= 2 or set(para.replace(" ", "")) <= {"-", "|", ":"}:
                    continue
                if para.lstrip().startswith(("-", "*", "•")):
                    doc.add_paragraph(para.lstrip("-*• ").strip(), style="List Bullet")
                else:
                    doc.add_paragraph(para)
        if section.table_data:
            _add_table(doc, section.table_data)


def render_document(plan: ExecutionPlan, user_request: str) -> str:
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()
    _set_base_styles(doc)
    _add_footer(doc)
    _add_title_page(doc, plan, user_request)
    _add_assumptions(doc, plan)
    _add_sections(doc, plan)

    safe_title = "".join(c if c.isalnum() or c in " _-" else "" for c in plan.title).strip().replace(" ", "_")
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{safe_title or 'document'}_{timestamp}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath
